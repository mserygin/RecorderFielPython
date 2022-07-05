import json
import os
import sys
from flask import Flask, request, jsonify
from docx import Document
import convertapi
import fitz
from flask_cors import CORS, cross_origin

app = Flask(__name__)
CORS(app, origins="*")

absPath = os.path.dirname(sys.modules['__main__'].__file__)
INPUT_FILE_PATH = absPath + "/myfolder/document.docx"
OUTPUT_FILE_PATH = absPath + "/myfolder/document.pdf"
CONVERT_KEY = "G3LHidhmyZOOYX56"


@app.route('/check-text/', methods=['POST'])
@cross_origin()
def CheckDocument():
    file = request.files['file']
    paramsText = json.loads(request.form.get('data'))

    file.save(INPUT_FILE_PATH)
    convertapi.api_secret = CONVERT_KEY
    result = convertapi.convert('pdf', {'File': INPUT_FILE_PATH})
    result.file.save(OUTPUT_FILE_PATH)

    try:
        return jsonify(checkDocument(file, paramsText))
    except ValueError:
        return ValueError


def checkDocument(dataFile, params):
    atr = []
    document = Document(dataFile)
    resultTest = {}
    result = []

    for index, paragraph in enumerate(document.paragraphs):
        for param in params:
            if param == 'fontSize':
                atr = atr + getFontSizeForParagraph(paragraph, document)
            if param == 'fontName':
                atr = atr + getFontsForParagraph(paragraph, document)
            if param == 'firstLineIndent':
                atr = atr + getFirstLineIndentForParagraph(paragraph, document)

            unique = list(map(lambda x: str(x), list(set(atr))))
            if unique != [params.get(param)]:
                if param in resultTest:
                    resultTest[param] = resultTest[param] + [getPageByParagraph(paragraph.text)]
                else:
                    resultTest[param] = [getPageByParagraph(paragraph.text)]
            atr = []
    for res in resultTest:
        if len(resultTest[res]) != 1 and resultTest[res] is not None:
            result.append({'name': res, 'listProblemsPage': list(filter(None, list(set(resultTest[res]))))})
    return result


def getPageByParagraph(paragraph):
    rim = fitz.open(OUTPUT_FILE_PATH)
    for current_page in range(len(rim)):
        page = rim.load_page(current_page)
        if page.search_for(paragraph):
            return current_page + 1


def getFontSizeForParagraph(paragraph, document):
    result = []
    for run in paragraph.runs:
        fontSizeParagraph = run.font.size
        fontSizeFromStyles = document.styles[getStyleParagraph(paragraph)].font.size
        if fontSizeParagraph == 0:
            result.append(round(0))
        elif fontSizeParagraph:
            result.append(round(fontSizeParagraph.pt))
        else:
            result.append(round(fontSizeFromStyles.pt))
    return result


def getFirstLineIndentForParagraph(paragraph, document):
    lineIndentParagraph = paragraph.paragraph_format.first_line_indent
    lineIndentFromStyles = document.styles[getStyleParagraph(paragraph)].paragraph_format.first_line_indent
    if lineIndentParagraph == 0:
        return [round(0, 2)]
    elif lineIndentParagraph:
        return [round(lineIndentParagraph.cm, 2)]
    return [round(lineIndentFromStyles.cm, 2)]


def getFontsForParagraph(paragraph, document):
    result = []
    for run in paragraph.runs:
        result.append(run.font.name or document.styles[getStyleParagraph(paragraph)].font.name)
    return result


def getStyleParagraph(paragraph): return paragraph.style.name


@app.route('/', methods=['GET'])
@cross_origin()
def getStatusProject():
    return "Everything works!"


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
